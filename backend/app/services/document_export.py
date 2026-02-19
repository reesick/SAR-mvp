from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from typing import Dict

class DocumentExporter:
    """Export SAR to Word and PDF"""
    
    def export_word(self, sar_data: Dict) -> io.BytesIO:
        """
        Export SAR to Word document with hyperlinks.
        
        Args:
            sar_data: Dict with sar_draft, risk_score, case_id, audit_data
            
        Returns:
            BytesIO buffer with Word document
        """
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Suspicious Activity Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        doc.add_heading('Case Information', level=1)
        doc.add_paragraph(f"Case ID: {sar_data['case_id']}")
        doc.add_paragraph(f"Risk Score: {sar_data['risk_score']}")
        doc.add_paragraph(f"Recommendation: {sar_data['recommendation']}")
        
        # SAR Narrative
        doc.add_heading('Narrative', level=1)
        
        # Split into paragraphs and add hyperlinks
        paragraphs = sar_data['sar_draft'].split('\n\n')
        
        for idx, para_text in enumerate(paragraphs):
            if para_text.strip():
                p = doc.add_paragraph(para_text.strip())
                
                # Add hyperlink to audit data - simplified as marker for now
                run = p.add_run(f" [Ref #{idx+1}]")
                run.font.color.rgb = None # default
                run.font.italic = True
        
        # Audit Appendix
        doc.add_page_break()
        doc.add_heading('Audit Trail', level=1)
        
        if 'audit_logs' in sar_data:
            for log in sar_data['audit_logs']:
                doc.add_paragraph(
                    f"{log['agent_name']} - {log['action_type']} at {log['timestamp']}",
                    style='List Bullet'
                )
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return buffer
    
    def export_pdf(self, sar_data: Dict) -> io.BytesIO:
        """
        Export SAR to PDF using xhtml2pdf.
        
        Args:
            sar_data: Dict with sar_draft, risk_score, case_id
            
        Returns:
            BytesIO buffer with PDF
        """
        
        from xhtml2pdf import pisa
        
        # Build HTML (reused logic)
        html_content = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial; margin: 40px; }}
                h1 {{ color: #333; text-align: center; }}
                .metadata {{ background: #f0f0f0; padding: 15px; margin: 20px 0; }}
                .narrative {{ line-height: 1.6; }}
                .audit {{ margin-top: 30px; border-top: 2px solid #ccc; padding-top: 20px; }}
            </style>
        </head>
        <body>
            <h1>Suspicious Activity Report</h1>
            
            <div class="metadata">
                <p><strong>Case ID:</strong> {sar_data['case_id']}</p>
                <p><strong>Risk Score:</strong> {sar_data['risk_score']}</p>
                <p><strong>Recommendation:</strong> {sar_data['recommendation']}</p>
            </div>
            
            <h2>Narrative</h2>
            <div class="narrative">
                {sar_data['sar_draft'].replace(chr(10), '<br>')}
            </div>
            
            <div class="audit">
                <h2>Audit Trail</h2>
                <ul>
        """
        
        if 'audit_logs' in sar_data:
            for log in sar_data['audit_logs']:
                html_content += f"<li>{log['agent_name']} - {log['action_type']} at {log['timestamp']}</li>"
                
        html_content += """
                </ul>
            </div>
        </body>
        </html>
        """
        
        # Convert to PDF
        buffer = io.BytesIO()
        pisa_status = pisa.CreatePDF(html_content, dest=buffer)
        
        if pisa_status.err:
            raise Exception("PDF generation error")
            
        buffer.seek(0)
        
        return buffer
